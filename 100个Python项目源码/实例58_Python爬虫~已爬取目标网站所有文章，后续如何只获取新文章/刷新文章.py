print("文章刷新中......")

#定义函数，获取想要的文章并批量写入word文件
import requests
from bs4 import BeautifulSoup
import docx
from docx.shared import Pt #用于设定字体大小（磅值）
from docx.oxml.ns import qn #用于应用中文字体
import random
import time
import os

def Get_article_to_word(url,date):
    user_agent_list = ["Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/68.0.3440.106 Safari/537.36",
                    "Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/67.0.3396.99 Safari/537.36",
                    "Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/64.0.3282.186 Safari/537.36",
                    "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/62.0.3202.62 Safari/537.36",
                    "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/45.0.2454.101 Safari/537.36",
                    "Mozilla/4.0 (compatible; MSIE 7.0; Windows NT 6.0)",
                    "Mozilla/5.0 (Macintosh; U; PPC Mac OS X 10.5; en-US; rv:1.9.2.15) Gecko/20110303 Firefox/3.6.15"
                    ]
    header = {'User-Agent': user_agent_list}
    header['User-Agent'] = random.choice(user_agent_list) #每篇文章随机选择浏览器，避免单个浏览器请求太快被服务器切断连接
    wb_data = requests.get(url,headers = header)
    soup = BeautifulSoup(wb_data.content)
    title = soup.select('.headword')[0].text.strip()#获得标题
    content1 = soup.select(".MsoNormal") #针对正文布局为 class = "MsoNormal"
    content2 = soup.select("#art_content") #针对正文布局为 id = "art_content"

    doc = docx.Document() #新建空白word文档
    #设定全局字体
    doc.styles['Normal'].font.name=u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    #写入标题行，并设置字体格式
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(18)

    doc.add_paragraph(date) #写入日期
    doc.add_paragraph(url) #写入文章链接

    dirs = os.getcwd() + "\\文章"
    if not os.path.exists(dirs):
        os.makedirs(dirs)
    #写入正文
    for i in content2:
        doc.add_paragraph(i.text)
    for i in content1:
        doc.add_paragraph(i.text)
    doc.save(f"{dirs}\\{title}.docx")


#判断网站是否有新文章
import requests
import json

header = {'User-Agent': 'Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/68.0.3440.75 Safari/537.36'}
form_data = {'_q': 'Article.list',
'siteId': '7e0b3b27-2622-4aa7-b6f8-abfe5c5df922',
'catalogId': '34f92da3-d6d0-4e96-899f-d7f581c18162',
'pub': 'true',
'limit': 150, #网站更新速度较慢，每次查看是否有更新的时候，只获取前10页共150篇文章的信息
'start': 1}

#这是异步加载，请求方法是POST
url = "http://www.bicpa.org.cn/dtzj/zxgg/getArticles.action"
res = requests.post(url, data = form_data, headers = header)
article_data0 =res.text.split("{success:true,datas:")[1] #去掉字符串前面的无用信息“{success:true,datas:”
article_data = article_data0.split(",total:")[0] #去掉字符串后面的无用信息“,total:xxxx}”
obj = json.loads(article_data)

#载入数据库中的所有链接，作为判断基准
f = open(os.getcwd()+'\\links.txt', 'r')
link_database = f.read() #读取成一个大字符串

#将所有新文章链接写入文本文件
path = r"http://www.bicpa.org.cn"
links_file = open('links.txt', 'a') #避免覆盖已有数据，用添加模式`a`写入
counter_link = 0 #新文章链接计数器
counter_download = 0 #新下载文章计数器
for info in obj:
    link = path + info['url'] + info['primaryKey']+".html" #拼接链接信息
    if link in link_database: #判断新提取的链接是否已存在于数据库
        pass
    else:
        links_file.write(link+"\n") #写入链接信息
        counter_link +=1 

        #下载标题含有"委员会专家提示"的文章
        if "委员会专家提示" in info['title']:
            Get_article_to_word(link,info["publishDate"]) #调用写好的函数，下载文章到word文件
            counter_download += 1       #每下载一篇文章，计数器增加1

links_file.close() #写完后关闭文件

#显示每次的结果
if counter_link == 0:
    print("没有文章更新！")
else:
    print(f"共获取到 {counter_link} 篇新文章的链接，并加入数据库。")

if counter_download == 0:
    print("没有'委员会专家提示'文章更新！")
else:
    print(f"共下载 {counter_download} 篇'委员会专家提示'新文章，请到文件夹查看。")

print("程序运行完成，关闭窗口退出.")
input()