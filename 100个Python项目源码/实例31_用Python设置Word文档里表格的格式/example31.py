#单元格边框设置函数
from docx.table import _Cell
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
 
def Set_cell_border(cell: _Cell, **kwargs):
    """
    设置单元格边框函数
    使用方法:
    Set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    传入参数有cell, 即单元格；top指上边框；bottom指下边框；start指左边框；end指右边框。
    "sz"指线的粗细程度；"val"指线型，比如单线，虚线等；"color"指颜色，颜色编码可百度；
    "space"指间隔，一般不设置，设置的值大于0会导致线错开；"shadow"指边框阴影
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
 
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

#定义单元格填充颜色函数
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
def Set_Background_Color(cell,rgbColor):
    shading_elm = parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'),color_value = rgbColor)) #固定写法，照抄即可
    cell._tc.get_or_add_tcPr().append(shading_elm)

from docx.enum.text import WD_ALIGN_PARAGRAPH #设置水平居中对齐需要用到的库
from docx.enum.table import WD_ALIGN_VERTICAL #设置垂直居中对齐需要用到的库
from docx.shared import Cm # 长度单位（厘米cm用到的库)
from docx import Document

doc = Document("收货记录.docx") #打开word文件
table= doc.tables[0]
max_row = len(table.rows) #获取表格总行数

#1.最后一行中，“总数”两个字要加粗
run = table.cell(max_row-1,4).paragraphs[0].runs[0] # 获取最后一行第五列对应单元格中的文字块，即“总数”字样所在单元格
run.font.bold = True #将文字块设置为粗体

#2.最后一行的行高要调大一点，现在这样扁扁的，难看
table.rows[max_row-1].height = Cm(1) # 将最后一行的行高设置为1厘米

#3.最后一行那几个空单元格有框线，太碍眼，得去掉，将总数及左边两个cell左边框和下边框设成白色
for i in range(3):
    cell = table.cell(max_row-1,i)
    Set_cell_border(
            cell,
            bottom={"color": "#FFFFFF"},
            start={"color": "#FFFFFF" },
            end={"color": "#FFFFFF"}
        )

cell_1 = table.cell(max_row-1,3)
Set_cell_border(cell_1,start={"color": "#FFFFFF" },bottom={"color": "#FFFFFF"})

cell_2 = table.cell(max_row-1,6)
Set_cell_border(cell_2,end={"color": "#FFFFFF" },bottom={"color": "#FFFFFF"})

#4.各单元格水平和垂直方向都要居中对齐
for row in range(1,max_row):
    for col in range(len(table.columns)):
        table.cell(row,col).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(row,col).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
#5.数量列≥85的，底色设置为橙色
qty = [] #存储数量信息
#读取第二行到29行，第2，3列中的数据
for i in range(1,max_row-1):
    qty_info = table.rows[i].cells[5].text #cells[5]指表格第6列
    qty.append(int(qty_info))
    
#将数量≥85的单元格填色
row=1 #行计数器
for i in qty:
    if i>=85:
        cell = table.cell(row,5) #第6列数据为数量，列索引是5
        Set_Background_Color(cell,"98F5FF") #填充颜色，"98F5FF"是蓝色的编码
    row+=1 #跳转到下一行      

doc.save("收货记录-整理.docx")