print("开始爬取文章......")
import requests
import os
import json

header = {'User-Agent': 'Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/68.0.3440.75 Safari/537.36'}

form_data = {'_q': 'Article.list',
'siteId': '7e0b3b27-2622-4aa7-b6f8-abfe5c5df922',
'catalogId': '34f92da3-d6d0-4e96-899f-d7f581c18162',
'pub': 'true',
'limit': 5000,
'start': 1}

#这是异步加载，请求方法是POST
url = "http://www.bicpa.org.cn/dtzj/zxgg/getArticles.action"
res = requests.post(url, data = form_data, headers = header)

article_data0 =res.text.split("{success:true,datas:")[1] #去掉字符串前面的无用信息“{success:true,datas:”
article_data = article_data0.split(",total:")[0] #去掉字符串后面的无用信息“,total:4946}”

obj = json.loads(article_data)
#获取标题含有“委员会专家提示”的文章的标题，发布时间和链接
path = r"http://www.bicpa.org.cn"
articles = []
for info in obj:
    if "委员会专家提示" in info['title']:
        article = {
            "标题": info['title'].strip(), #strip()去除首尾空格
            "发布时间": info['publishDate'],
            "链接": path + info['url'] + info['primaryKey']+".html"             
        }
        articles.append(article)

#获取想要的文章并批量写入word文件
import requests
from bs4 import BeautifulSoup
import docx
from docx.shared import Pt #用于设定字体大小（磅值）
from docx.oxml.ns import qn #用于应用中文字体
import random
import time

def Get_article_to_word(url,date):
    user_agent_list = ["Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/68.0.3440.106 Safari/537.36",
                    "Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/67.0.3396.99 Safari/537.36",
                    "Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/64.0.3282.186 Safari/537.36",
                    "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/62.0.3202.62 Safari/537.36",
                    "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/45.0.2454.101 Safari/537.36",
                    "Mozilla/4.0 (compatible; MSIE 7.0; Windows NT 6.0)",
                    "Mozilla/5.0 (Macintosh; U; PPC Mac OS X 10.5; en-US; rv:1.9.2.15) Gecko/20110303 Firefox/3.6.15"
                    ]
    header = {'User-Agent': user_agent_list}
    header['User-Agent'] = random.choice(user_agent_list) #每篇文章随机选择浏览器，避免单个浏览器请求太快被服务器切断连接
    wb_data = requests.get(url,headers = header)
    soup = BeautifulSoup(wb_data.content)
    title = soup.select('.headword')[0].text.strip()#获得标题
    content1 = soup.select(".MsoNormal") #针对正文布局为 class = "MsoNormal"
    content2 = soup.select("#art_content") #针对正文布局为 id = "art_content"
    
    doc = docx.Document() #新建空白word文档
    #设定全局字体
    doc.styles['Normal'].font.name=u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
   
    #写入标题行，并设置字体格式
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(18)
    
    doc.add_paragraph(date) #写入日期
    doc.add_paragraph(url) #写入文章链接
    
    dirs = os.getcwd() + "\\文章"
    if not os.path.exists(dirs):
        os.makedirs(dirs)

    #写入正文
    for i in content2:
        doc.add_paragraph(i.text)
    for i in content1:
        doc.add_paragraph(i.text)
    doc.save(f"{dirs}\\{title}.docx")

#遍历所有文章的链接，调用以上函数执行
for art in articles:
    Get_article_to_word(art["链接"],art["发布时间"])
    print("{} 下载完成。".format(art['标题']))
    if articles.index(art) % 30 == 29: #每获取30篇文章，暂停5秒，避免频繁请求被服务器切断连接
        time.sleep(5)
          
print(f"共下载 {len(articles)} 篇文章。")
print("程序运行完成，关闭窗口退出.")
input()