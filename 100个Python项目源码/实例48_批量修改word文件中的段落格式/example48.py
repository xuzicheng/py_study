#定义字体格式
from docx.shared import Pt #用于设定字体大小（磅值）
from docx.oxml.ns import qn #用于应用中文字体

def F_title(run):
    #标题文字    
    run.font.size = Pt(22) #文字大小磅值
    run.bold = True #加粗
    run.font.name = "方正小标宋_GBK" #字体
    r = run._element.rPr.rFonts
    r.set(qn("w:eastAsia"),"方正小标宋_GBK") #字体
        
def F_name_dept(run):
    #姓名，部门，日期
    run.font.size = Pt(17) #文字大小磅值
    run.bold = False #加粗
    run.font.name = "楷体" #字体
    r = run._element.rPr.rFonts
    r.set(qn("w:eastAsia"),"楷体") #字体
    
def F_main(run):        
    #正文的格式
    run.font.size = Pt(17) #文字大小磅值
    run.bold = False #加粗
    run.font.name = "仿宋" #字体
    r = run._element.rPr.rFonts
    r.set(qn("w:eastAsia"),"仿宋") #字体
    
def F_title1(run):
    #标题一的格式
    run.font.size = Pt(17) #文字大小磅值
    run.bold = False #加粗
    run.font.name = "黑体" #字体
    r = run._element.rPr.rFonts
    r.set(qn("w:eastAsia"),"黑体") #字体
    
def F_title2(run):
    #标题二的格式
    run.font.size = Pt(17) #文字大小磅值
    run.bold = True #加粗
    run.font.name = "楷体" #字体
    r = run._element.rPr.rFonts
    r.set(qn("w:eastAsia"),"楷体") #字体

import docx,os
#获取待处理的文件的路径
path='待处理文件'  #文件所在文件夹
files = [path+"\\"+i for i in os.listdir(path)] #获取文件夹下的文件名,并拼接完整路径

#逐个提取文件，设置字体格式
for file in files:
    doc = docx.Document(file)
    for run in doc.paragraphs[0].runs: #总标题字体格式
        F_title(run)
        
    for para in doc.paragraphs[1:3]: #部门、姓名及日期字体格式
        for run in para.runs:
            F_name_dept(run)
    
    title1 = ["一、","二、","三、","四、"] #标题一的唯一特征字符串
    title2 = ["1、","2、","3、","4、"] #标题二的唯一特征字符串
    for para in doc.paragraphs[3:]:
        if any(i in para.text for i in title1): #若该段落是标题一，则应用标题一的字体格式
            for run in para.runs:
                F_title1(run)
        elif any(j in para.text for j in title2):#若该段落是标题二，则应用标题二的字体格式
            for run in para.runs:
                F_title2(run)
        else:
            for run in para.runs: #其余都应用正文的字体格式
                F_main(run)
    doc.save('已处理文件\\{}'.format(file.split("\\")[1]))
        